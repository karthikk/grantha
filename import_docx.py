#!/usr/bin/env python3
"""Import texts from BVM Granthaah.docx into the grantha site."""

import os
import re
from docx import Document
from build import *

DOCX_PATH = os.path.expanduser("~/Downloads/BVM Granthaah.docx")
BASE = os.path.dirname(os.path.abspath(__file__))

# Texts already on the site — skip these
SKIP = {
    'तोटकाष्टकम्', 'गोविन्दाष्टकम्', 'द्वादश ज्योतिर्लिङ्ग स्तोत्रं',
    'भज गोविन्दं', 'धन्याष्टकम्', 'प्रातःस्मरणस्तोत्रम्',
    'काशीपञ्चकम्', 'दशश्लोकी', 'मनीषापञ्चकम्', 'मायापञ्चकम्',
    'श्रीमद्भगवद्गीता',
    # Non-stotra items
    'वन्दे मातरम्', 'जन्मदिन गीतम्', 'विवाहदिनन गीतम्',
}

# Classify as prakarana or sloka
PRAKARANA_TITLES = {
    'प्रश्नोत्तररत्नमालिका', 'तत्त्वबोधः', 'आत्मबोधः',
    'लघुवाक्यवृत्तिः', 'वाक्यवृत्तिः', 'एकश्लोकी',
    'अद्वैतपञ्चरत्नम्', 'हस्तामलकीयम्', 'अद्वैतानुभूतिः',
    'ब्रह्मज्ञानावलीमाला', 'साधन पञ्चकम्',
}

# Title -> (clean_title, filename_key)
TITLE_MAP = {
    'गुरु स्तोत्रम्': ('गुरुस्तोत्रम्', 'gurustotram'),
    'गुरु पादुका स्तोत्रम्': ('गुरुपादुकास्तोत्रम्', 'gurupadukastotram'),
    'गुरु अष्टकम्': ('गुर्वष्टकम्', 'gurvashtakam'),
    'निर्वाण षट्कम्': ('निर्वाणषट्कम्', 'nirvanashatkam'),
    'संकट नाशन गणेश स्तोत्रं': ('सङ्कटनाशनगणेशस्तोत्रम्', 'sankatanashana'),
    'दक्षिणामूर्त्यष्टकम्': ('दक्षिणामूर्त्यष्टकम्', 'dakshinamurtyashtakam'),
    'शिव मनसा पूजा': ('शिवमानसपूजा', 'shivamanasapuja'),
    'लिङ्गाष्टकम्': ('लिङ्गाष्टकम्', 'lingashtakam'),
    'मार्गबन्धुस्तोत्रम्': ('मार्गबन्धुस्तोत्रम्', 'margabandhustotram'),
    'नटराजस्तोत्रम्': ('नटराजस्तोत्रम्', 'natarajastotram'),
    'शिवशङ्कर स्तोत्रम्': ('शिवशङ्करस्तोत्रम्', 'shivashankarastotram'),
    'कृष्णाष्टकम्': ('कृष्णाष्टकम्', 'krishnashtakam'),
    'राम भुजङ्ग स्तोत्रम्': ('रामभुजङ्गस्तोत्रम्', 'ramabhujangastotram'),
    'गंगा स्तोत्रम्': ('गङ्गास्तोत्रम्', 'gangastotram'),
    'कमलजदयिताष्टकम्': ('कमलजदयिताष्टकम्', 'kamalajadayitashtakam'),
    'महिषासुरमर्दिनि स्तोत्रम्': ('महिषासुरमर्दिनिस्तोत्रम्', 'mahishasuramardini'),
    'एकात्मता स्तोत्रम्': ('एकात्मतास्तोत्रम्', 'ekatmatastotram'),
    'स्वरूपानुसंधानाष्टकम्': ('स्वरूपानुसन्धानाष्टकम्', 'svarupanusandhana'),
    'प्रश्नोत्तररत्नमालिका': ('प्रश्नोत्तररत्नमालिका', 'prashnottararatnamalika'),
    'तत्त्वबोधः': ('तत्त्वबोधः', 'tattvabodha'),
    'आत्मबोधः': ('आत्मबोधः', 'atmabodha'),
    'लघुवाक्यवृत्तिः': ('लघुवाक्यवृत्तिः', 'laghuvakyavritti'),
    'वाक्यवृत्तिः': ('वाक्यवृत्तिः', 'vakyavritti'),
    'एकश्लोकी': ('एकश्लोकी', 'ekashloki'),
    'अद्वैतपञ्चरत्नम्': ('अद्वैतपञ्चरत्नम्', 'advaitapancharatnam'),
    'हस्तामलकीयम्': ('हस्तामलकीयम्', 'hastamalakiyam'),
    'अद्वैतानुभूतिः': ('अद्वैतानुभूतिः', 'advaitanubhuti'),
    'ब्रह्मज्ञानावलीमाला': ('ब्रह्मज्ञानावलीमाला', 'brahmajnanavali'),
    'साधन पञ्चकम्': ('साधनपञ्चकम्', 'sadhanapanchakam'),
    'द्वादश ज्योतिर्लिङ्ग स्तोत्रं': ('', ''),  # skip
}


def fix_text(text):
    """Auto-fix common issues."""
    # Pipe to danda
    text = text.replace('|', '।')
    # || to double danda
    text = re.sub(r'॥\s*॥', '॥', text)
    text = re.sub(r'।।', '॥', text)
    # English S used as avagraha
    text = text.replace('S', 'ऽ')
    # Strip parenthetical notes like (इन्द्रवज्रा छन्द -)
    text = re.sub(r'\([^)]*छन्द[^)]*\)', '', text)
    # Strip inline corrections like (गोष्टी)
    text = re.sub(r'\([^)]*\)', '', text)
    # Normalize spaces
    text = re.sub(r'  +', ' ', text)
    text = text.strip()
    return text


def extract_sections(doc):
    """Extract all H1 sections from docx."""
    sections = []
    current_title = None
    current_lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ''

        if 'Heading 1' in style:
            if current_title:
                sections.append((current_title, current_lines))
            current_title = text
            current_lines = []
        elif 'Heading 2' in style:
            current_lines.append(('H2', text))
        else:
            if current_title and re.search(r'[\u0900-\u097F]', text):
                current_lines.append(('V', fix_text(text)))

    if current_title:
        sections.append((current_title, current_lines))

    return sections


def lines_to_verse_html(lines):
    """Convert list of (type, text) lines to verse HTML."""
    body = ''
    current_verse = []

    for ltype, text in lines:
        if ltype == 'H2':
            continue
        current_verse.append(text)
        # Check if this line ends a verse (has ॥ N ॥)
        if re.search(r'॥\s*[\d०-९]+\s*॥', text):
            verse_text = '\n'.join(current_verse)
            # Extract verse number
            m = re.search(r'॥\s*([\d०-९]+)\s*॥\s*$', verse_text)
            if m:
                num_str = m.group(1)
                # Convert to devanagari if needed
                if re.match(r'\d+', num_str):
                    num_str = to_dev(int(num_str))
                verse_text = verse_text[:m.start()].rstrip()
                vlines = verse_text.split('\n')
                inner = '<br>\n      '.join(l.strip() for l in vlines if l.strip())
                body += f'  <div class="verse">\n    <div class="shloka">\n      {inner} <span class="verse-num">॥ {num_str} ॥</span>\n    </div>\n  </div>\n'
            else:
                vlines = verse_text.split('\n')
                inner = '<br>\n      '.join(l.strip() for l in vlines if l.strip())
                body += f'  <div class="verse">\n    <div class="shloka">\n      {inner}\n    </div>\n  </div>\n'
            current_verse = []

    # Remaining lines without verse number
    if current_verse:
        text = '\n'.join(current_verse)
        vlines = text.split('\n')
        inner = '<br>\n      '.join(l.strip() for l in vlines if l.strip())
        if inner:
            body += f'  <div class="verse">\n    <div class="shloka">\n      {inner}\n    </div>\n  </div>\n'

    return body


def extract_gita_dhyanam(doc):
    """Extract Gita dhyana slokas."""
    lines = []
    in_dhyanam = False
    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name if para.style else ''
        if 'Heading 2' in style and 'ध्यानम्' in text:
            in_dhyanam = True
            continue
        if in_dhyanam:
            if 'Heading' in style:
                break
            if text and re.search(r'[\u0900-\u097F]', text):
                lines.append(('V', fix_text(text)))
    return lines


def main():
    doc = Document(DOCX_PATH)
    sections = extract_sections(doc)

    new_slokas = []
    new_prakaranas = []

    for title, lines in sections:
        if title in SKIP:
            continue
        if title not in TITLE_MAP:
            print(f"  Skipping unknown: {title}")
            continue

        clean_title, key = TITLE_MAP[title]
        if not key:
            continue

        is_prakarana = title in PRAKARANA_TITLES
        body = lines_to_verse_html(lines)

        if is_prakarana:
            new_prakaranas.append((clean_title, key, body))
        else:
            new_slokas.append((clean_title, key, body))

    # ─── Write new sloka pages ───
    sloka_dir = os.path.join(BASE, 'slokas')
    # Read existing slokas from current index
    existing_slokas = [
        ('तोटकाष्टकम्', 'totakashtakam'),
        ('द्वादशज्योतिर्लिङ्गस्तोत्रम्', 'dvadashalingastotram'),
        ('श्रीगोविन्दाष्टकम्', 'govindashtakam'),
    ]

    # Guru stotram should come first
    ordered_new_slokas = sorted(new_slokas, key=lambda x: x[1])
    # Put guru stotram first
    guru_first = [s for s in ordered_new_slokas if 'guru' in s[1].lower()]
    rest = [s for s in ordered_new_slokas if 'guru' not in s[1].lower()]
    ordered_new_slokas = guru_first + rest

    all_slokas = existing_slokas.copy()
    for clean_title, key, body in ordered_new_slokas:
        all_slokas.append((clean_title, key))

    print(f"\nWriting {len(ordered_new_slokas)} new sloka pages...")
    for si, (clean_title, key, body) in enumerate(ordered_new_slokas):
        # Find position in all_slokas for prev/next
        idx = next(i for i, (t, k) in enumerate(all_slokas) if k == key)
        prev_link = (f'{all_slokas[idx-1][1]}.html', all_slokas[idx-1][0]) if idx > 0 else None
        next_link = (f'{all_slokas[idx+1][1]}.html', all_slokas[idx+1][0]) if idx < len(all_slokas)-1 else None

        write_file(os.path.join(sloka_dir, f'{key}.html'),
                   page_html(clean_title,
                            [('../', 'मुख्यम्'), ('./', 'श्लोकाः'), (None, clean_title)],
                            body, prev_link=prev_link, next_link=next_link,
                            css_path='../css/style.css'))

    # Update slokas index
    items = ''
    for title, key in all_slokas:
        items += f'    <li><a href="{key}.html">{title}</a></li>\n'
    idx_body = f'  <ul class="text-list">\n{items}  </ul>\n'
    write_file(os.path.join(sloka_dir, 'index.html'),
               page_html('श्लोकाः',
                        [('../', 'मुख्यम्'), (None, 'श्लोकाः')],
                        idx_body, css_path='../css/style.css'))

    # ─── Write new prakarana pages ───
    prakarana_dir = os.path.join(BASE, 'prakarana')
    existing_prakaranas = [
        ('मायापञ्चकम्', 'mayapanchakam'),
        ('मनीषापञ्चकम्', 'manishapanchakam'),
        ('काशीपञ्चकम्', 'kashipanchakam'),
        ('दशश्लोकी', 'dashashloki'),
        ('धन्याष्टकम्', 'dhanyashtakam'),
        ('प्रातःस्मरणम्', 'pratahsmaranam'),
        ('भजगोविन्दम्', 'bhajagovindam'),
    ]

    all_prakaranas = existing_prakaranas.copy()
    for clean_title, key, body in new_prakaranas:
        all_prakaranas.append((clean_title, key))

    print(f"\nWriting {len(new_prakaranas)} new prakarana pages...")
    for pi, (clean_title, key, body) in enumerate(new_prakaranas):
        idx = next(i for i, (t, k) in enumerate(all_prakaranas) if k == key)
        prev_link = (f'{all_prakaranas[idx-1][1]}.html', all_prakaranas[idx-1][0]) if idx > 0 else None
        next_link = (f'{all_prakaranas[idx+1][1]}.html', all_prakaranas[idx+1][0]) if idx < len(all_prakaranas)-1 else None

        write_file(os.path.join(prakarana_dir, f'{key}.html'),
                   page_html(clean_title,
                            [('../', 'मुख्यम्'), ('./', 'प्रकरणग्रन्थाः'), (None, clean_title)],
                            body, prev_link=prev_link, next_link=next_link,
                            css_path='../css/style.css'))

    # Update prakarana index
    items = ''
    for title, key in all_prakaranas:
        items += f'    <li><a href="{key}.html">{title}</a></li>\n'
    idx_body = f'  <ul class="text-list">\n{items}  </ul>\n'
    write_file(os.path.join(prakarana_dir, 'index.html'),
               page_html('प्रकरणग्रन्थाः',
                        [('../', 'मुख्यम्'), (None, 'प्रकरणग्रन्थाः')],
                        idx_body, css_path='../css/style.css'))

    # ─── Gita dhyana slokas ───
    print("\nExtracting Gita dhyana slokas...")
    dhyanam_lines = extract_gita_dhyanam(doc)
    dhyanam_body = lines_to_verse_html(dhyanam_lines)
    # Prepend to Gita chapter 1
    gita_ch1_path = os.path.join(BASE, 'gita', 'adhyaya-1.html')
    with open(gita_ch1_path, 'r') as f:
        ch1_content = f.read()

    # Insert dhyanam section before the first verse
    dhyanam_section = f'\n  <h2 class="section-heading" id="dhyanam">ध्यानम्</h2>\n\n{dhyanam_body}\n  <h2 class="section-heading" id="adhyaya">अध्यायः</h2>\n\n'

    # Insert after <h1>
    ch1_content = re.sub(
        r'(</h1>\n)',
        r'\1' + dhyanam_section,
        ch1_content,
        count=1
    )
    write_file(gita_ch1_path, ch1_content)

    print("\nDone!")
    print(f"  New slokas: {len(ordered_new_slokas)}")
    print(f"  New prakaranas: {len(new_prakaranas)}")
    print(f"  Gita dhyanam added")


if __name__ == '__main__':
    main()
